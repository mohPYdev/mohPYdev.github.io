#!/usr/bin/env python3
"""Generate two tailored resumes (academic + industry) as .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ACCENT = RGBColor(0x8B, 0x5A, 0x2B)   # warm brown, matches the site
DARK   = RGBColor(0x22, 0x22, 0x22)
GREY   = RGBColor(0x55, 0x55, 0x55)
NAME   = "Mohammad Rafieian"
CONTACT = ("(945) 527-5229  |  mohammad.rafieian@utdallas.edu  |  mohpydev.github.io  |  "
           "linkedin.com/in/mohammad-rafieian  |  github.com/mohPYdev")

RIGHT_MARGIN_TAB = Inches(7.1)  # right edge of text (letter, 0.7in margins)


def setup(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.45)
        s.bottom_margin = Inches(0.45)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 0.95


def _no_space(p):
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space(p)
    r = p.add_run(NAME)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    r.font.name = 'Cambria'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(CONTACT)
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY


def section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    r.font.name = 'Cambria'
    # bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '8B5A2B')
    pbdr.append(bottom)
    pPr.append(pbdr)


def heading_line(doc, left_bold, left_rest, right, before=4):
    """A line with bold left text + normal continuation, and right-aligned date."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(0)
    # right tab stop
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_MARGIN_TAB, WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(left_bold)
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK
    if left_rest:
        r2 = p.add_run(left_rest)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
    if right:
        r3 = p.add_run('\t' + right)
        r3.font.size = Pt(9.1)
        r3.font.italic = True
        r3.font.color.rgb = GREY
    return p


def sub_line(doc, text, italic=True):
    p = doc.add_paragraph()
    _no_space(p)
    r = p.add_run(text)
    r.font.size = Pt(9.1)
    r.font.italic = italic
    r.font.color.rgb = GREY
    return p


def bullet(doc, segments):
    """segments: list of (text, bold) tuples, or a plain string."""
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    b = p.add_run('▪  ')
    b.font.color.rgb = ACCENT
    b.font.size = Pt(8)
    if isinstance(segments, str):
        segments = [(segments, False)]
    for text, bold in segments:
        r = p.add_run(text)
        r.font.size = Pt(9.1)
        r.font.bold = bold
        r.font.color.rgb = DARK
    return p


def pub_entry(doc, n, title, authors, venue):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    idx = p.add_run(f"{n}.  ")
    idx.font.size = Pt(9.1)
    idx.font.bold = True
    idx.font.color.rgb = ACCENT
    t = p.add_run(f"“{title}.” ")
    t.font.size = Pt(9.1)
    t.font.bold = True
    t.font.color.rgb = DARK
    # authors, bolding M. Rafieian
    parts = authors.split("M. Rafieian")
    for i, seg in enumerate(parts):
        if seg:
            r = p.add_run(seg)
            r.font.size = Pt(9.1)
            r.font.color.rgb = GREY
        if i < len(parts) - 1:
            rb = p.add_run("M. Rafieian")
            rb.font.size = Pt(9.1)
            rb.font.bold = True
            rb.font.color.rgb = DARK
    v = p.add_run(f" {venue}")
    v.font.size = Pt(9.1)
    v.font.italic = True
    v.font.color.rgb = ACCENT


def skills_line(doc, label, items):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.02)
    r = p.add_run(f"{label}:  ")
    r.font.bold = True
    r.font.size = Pt(9.1)
    r.font.color.rgb = DARK
    r2 = p.add_run(items)
    r2.font.size = Pt(9.1)
    r2.font.color.rgb = GREY


def summary(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(9.7)
    r.font.color.rgb = DARK


# ------------------------------------------------------------------ DATA

EDU = [
    ("University of Texas at Dallas", ", M.S. & Ph.D. in Software Engineering & Security",
     "2023 – Present", "M.S. May 2026 · Ph.D. expected May 2027 · Advisor: Dr. Shiyi Wei"),
    ("Isfahan University", ", B.S. in Computer Engineering",
     "Aug 2019 – Jun 2023", "ACM Student Chapter Member"),
]

# Publications ordered: accepted/published, under review, in submission
PUBS = [
    ("Is Call Graph Pruning Really Effective? An Empirical Re-evaluation",
     "M. Rafieian, V. Birsan, K. Katiyar, D. Zhong, S. Wei.", "ICSE 2026."),
    ("Improving ML-based Static Analysis Classification via Explainable AI",
     "S. Yerramreddy, M. Rafieian, S. Wei, A. Porter.", "ICST 2026."),
    ("Agentic Root Cause Reasoning and Unsoundness in Call Graph Analysis",
     "M. Rafieian, M. Hicks, S. Wei.", "ICSE 2027 (under review)."),
    ("Automated White-Box Unit Testing for Functional Requirements",
     "Y. Zhou, M. Rafieian, S. Wei, W. Lam, A. Marcus.", "ICSE 2027 (under review)."),
    ("An Empirical Study of Static Analysis-Based Variability Bug Detection",
     "A. Mordahl, Z. Patterson, M. Rafieian, M. T. Ahmad, S. Wei.", "ESEM 2026 (under review)."),
    ("Automatic Test Suites for Static Analysis Tools via Dynamic Analysis",
     "A. Mordahl, M. Rafieian, S. Wei.", "FSE 2027 (in submission)."),
]


def education(doc):
    section(doc, "Education")
    for bold, rest, date, detail in EDU:
        heading_line(doc, bold, rest, date, before=3)
        if detail:
            sub_line(doc, detail)


# ============================================================ ACADEMIC CV

def build_academic():
    doc = Document()
    setup(doc)
    header(doc)

    section(doc, "Research Interests")
    summary(doc,
        "Program analysis, automated software testing, and software security — applying machine learning and "
        "rigorous empirical methods to improve the soundness and precision of static analysis and call-graph tools.")

    education(doc)

    section(doc, "Publications")
    for i, (title, authors, venue) in enumerate(PUBS, 1):
        pub_entry(doc, i, title, authors, venue)

    section(doc, "Research Experience")
    heading_line(doc, "Graduate Research Assistant", " — University of Texas at Dallas",
                 "Aug 2023 – Present")
    sub_line(doc, "Systems Software & Security Lab · Advisor: Dr. Shiyi Wei")
    for seg in [
        "Lead research in program analysis and security, applying machine learning to automated vulnerability detection and testing for C/Java systems.",
        "Designed large-scale, reproducible experimental pipelines using fuzzing and dynamic bytecode instrumentation to benchmark four static analysis tools.",
        "Introduced three novel datasets and evaluation methodologies for static analysis, resulting in six papers across ICSE, ICST, ESEM, and FSE.",
    ]:
        bullet(doc, seg)

    heading_line(doc, "Research Assistant", " — Model-Driven SE Lab, Isfahan University",
                 "Jan 2022 – Jan 2023")
    for seg in [
        "Built an automatic reservation-system generator engine (MDSE) and co-authored a journal paper on automatically generating full web applications.",
    ]:
        bullet(doc, seg)

    section(doc, "Selected Research Projects")
    heading_line(doc, "Call Graph Pruning Framework", "  ·  Python, Java, Bash", "Feb – Jul 2025")
    for seg in [
        "Built static call graphs (DOOP, OPAL, WALA) and instrumented Java bytecode (ASM); fine-tuned CodeBERT/CodeT5 to improve edge classification by 30%.",
        "Re-evaluated prior pruning methods and fixed dataset flaws, showing reported gains were inflated (+50% claimed vs. +5% actual).",
    ]:
        bullet(doc, seg)
    heading_line(doc, "Explainable AI for Static Analysis", "  ·  Python, Java, C, Bash", "Jan – Sep 2025")
    for seg in [
        "Constructed the first explainable dataset for static-analysis ML (1,132 Java + 67K C samples), evaluating 7 code LLMs.",
        "Applied Explanation-Guided Learning to raise relevant-line focus from 35% to 65% and improve classification accuracy by 20%.",
    ]:
        bullet(doc, seg)
    heading_line(doc, "Fuzzing for Static Analysis Benchmarking", "  ·  Python, Java, C", "Jul 2025 – Present")
    for seg in [
        "Built a hybrid (directed + grey-box) fuzzing pipeline with a novel feedback mechanism that auto-generates static-analysis benchmarks, saving ~80 person-hours.",
    ]:
        bullet(doc, seg)

    section(doc, "Academic Service")
    bullet(doc, [("Artifact Evaluation Committee", True),
                 (" — PLDI, ASE, ISSTA, and ICST (2026)", False)])

    section(doc, "Technical Skills")
    skills_line(doc, "Program Analysis", "WALA, DOOP, Soot, OPAL, Jazzer, pointer & call-graph analysis, bytecode instrumentation (ASM)")
    skills_line(doc, "Machine Learning", "PyTorch, TensorFlow, scikit-learn, LLM fine-tuning (CodeBERT/CodeT5), Explainable AI")
    skills_line(doc, "Languages & Tools", "Python, Java, C/C++, Scala, JavaScript, SQL · Git, Docker, Linux, PostgreSQL, Redis")

    doc.save("resume/cv-academic.docx")
    print("wrote resume/cv-academic.docx")


# ============================================================ INDUSTRY CV

def build_industry():
    doc = Document()
    setup(doc)
    header(doc)

    section(doc, "Summary")
    summary(doc,
        "Software engineer and Ph.D. researcher who builds scalable, secure systems end-to-end. Strong in "
        "backend architecture, ML pipelines, and large-scale automation, with deep expertise in program "
        "analysis and software security. Experienced shipping Dockerized full-stack applications "
        "(Django/React) and designing reproducible, data-intensive infrastructure.")

    section(doc, "Technical Skills")
    skills_line(doc, "Languages", "Python, Java, C, C++, JavaScript, Scala, SQL")
    skills_line(doc, "Backend & Web", "Django, Flask, Node.js, React, REST APIs, WebSockets")
    skills_line(doc, "ML & Data", "PyTorch, TensorFlow, scikit-learn, LLM fine-tuning, ML pipelines")
    skills_line(doc, "Infrastructure", "Docker, Linux, Git, PostgreSQL, Redis, Firebase, Vercel")
    skills_line(doc, "Focus Areas", "Program Analysis, Automated Testing, Software Security, Backend Systems")

    section(doc, "Experience")
    heading_line(doc, "Graduate Research Assistant", " — University of Texas at Dallas",
                 "Aug 2023 – Present")
    for seg in [
        "Designed and maintained large-scale, reproducible experimental pipelines using fuzzing and dynamic instrumentation to benchmark four static analysis tools.",
        "Built ML training and evaluation infrastructure (PyTorch/TensorFlow) with automated hyperparameter tuning and cross-validation, fine-tuning transformer models to +30% accuracy.",
        "Containerized experiments with Docker on Linux for portability and reproducibility; authored six peer-reviewed papers.",
    ]:
        bullet(doc, seg)

    heading_line(doc, "Research Assistant", " — Model-Driven SE Lab, Isfahan University",
                 "Jan 2022 – Jan 2023")
    for seg in [
        "Developed an automatic reservation-system generator engine that produced full web applications from high-level models.",
    ]:
        bullet(doc, seg)

    section(doc, "Projects")
    heading_line(doc, "ArtCom — Online Art Auction Platform", "  ·  React, Django, PostgreSQL, Docker",
                 "Feb – Jun 2022")
    for seg in [
        "Built a real-time art auction and exhibition platform (React.js + Django REST Framework), cutting renting costs by 90%.",
        "Implemented a Dockerized, WebSocket-based bidding system supporting 200 concurrent participants with Redis-backed synchronization.",
    ]:
        bullet(doc, seg)
    heading_line(doc, "LORD — Low-Code Reservation Platform", "  ·  Django, React, MySQL", "Aug – Dec 2022")
    for seg in [
        "Pioneered a low-code reservation-system generator that automated website creation, cutting deployment costs by over 90%.",
        "Engineered a scalable full-stack solution with integrated payment processing for seamless deployment.",
    ]:
        bullet(doc, seg)
    heading_line(doc, "Call Graph Pruning Framework", "  ·  Python, Java, PyTorch", "Feb – Jul 2025")
    for seg in [
        "Fine-tuned CodeBERT/CodeT5 with custom training pipelines to improve call-graph edge classification accuracy by 30%.",
        "Re-evaluated prior methods and fixed dataset flaws, showing reported gains were inflated (+50% claimed vs. +5% actual).",
    ]:
        bullet(doc, seg)

    section(doc, "Education")
    for bold, rest, date, detail in EDU:
        heading_line(doc, bold, rest, date, before=3)
        if detail:
            sub_line(doc, detail)

    section(doc, "Research & Publications")
    summary(doc,
        "Author of six papers in top software-engineering venues (ICSE, ICST, ESEM, FSE) on program analysis, "
        "automated testing, and ML for security. Full list at mohpydev.github.io.")

    doc.save("resume/cv-industry.docx")
    print("wrote resume/cv-industry.docx")


if __name__ == "__main__":
    build_academic()
    build_industry()
